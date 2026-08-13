"""
Generate Weekly_GDB_Gap_Report.docx from analysis_output.json or pipeline output.
"""
import os
import json
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    Document = None

def generate_report():
    json_path = os.path.join(os.path.dirname(__file__), "analysis_output.json")
    if not os.path.exists(json_path):
        import analyze_gaps
        data = analyze_gaps.output
    else:
        with open(json_path) as f:
            data = json.load(f)

    if Document is None:
        print("docx library not installed. Install with `pip install python-docx` to generate Word doc.")
        return

    doc = Document()

    # Page Setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles & Colors
    # Primary: #2B3A67 (Indigo), Secondary: #B2412E (Rust), Neutral Dark: #241F1A
    def style_heading(p, text, size_pt=18, color_rgb=(43, 58, 103), bold=True):
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.font.color.rgb = RGBColor(*color_rgb)
        return run

    # Header / Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_after = Pt(2)
    r_sub = p_title.add_run("ACE · GDB KNOWLEDGE GAP INTELLIGENCE\n")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(9)
    r_sub.bold = True
    r_sub.font.color.rgb = RGBColor(92, 85, 70)

    style_heading(p_title, "Weekly GDB Coverage Gap Report", size_pt=24, color_rgb=(43, 58, 103))

    report = data["report"]
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(16)
    r_meta = p_meta.add_run(f"Report Week: {report['report_week']}  |  Period Covered: {report['period_covered']}  |  Generated: {datetime.utcnow().strftime('%Y-%m-%d')}")
    r_meta.font.name = "Arial"
    r_meta.font.size = Pt(9.5)
    r_meta.font.italic = True
    r_meta.font.color.rgb = RGBColor(100, 100, 100)

    # Executive Summary Box
    p_sec1 = doc.add_paragraph()
    style_heading(p_sec1, "1. Executive Summary", size_pt=14, color_rgb=(43, 58, 103))

    p_summary = doc.add_paragraph()
    p_summary.paragraph_format.space_after = Pt(12)
    p_summary.add_run(
        f"Over the 12-week analysis period, a total of {report['total_disclaimer_queries']:,} disclaimer-triggered queries "
        f"were collected across 10 agricultural domains and 10 major crop states. The clustering engine identified "
        f"{report['total_gap_clusters']} actionable knowledge gap clusters.\n\n"
        f"Key Highlights:\n"
        f"• Chronic High-Volume Gaps (≥ {report['high_volume_threshold']} queries/12wk): {len(data['high_volume_gaps'])} clusters require priority GDB content creation.\n"
        f"• Emerging Fast-Growing Spikes (≥ {report['fast_growth_threshold_pct']}% WoW growth): {len(data['fast_growing_gaps'])} clusters require rapid SOP/advisory deployment.\n"
        f"• Outreach Actions: {report['headline_stats']['immediate_actions_count']} immediate high-urgency field actions recommended."
    )

    # Section 2: High-Volume Gaps
    p_sec2 = doc.add_paragraph()
    style_heading(p_sec2, "2. High-Volume Knowledge Gaps (Chronic Blind Spots)", size_pt=14, color_rgb=(43, 58, 103))

    table_hv = doc.add_table(rows=1, cols=5)
    table_hv.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_hv.autofit = False

    hdr_cells = table_hv.rows[0].cells
    headers = ["Crop & State", "Domain", "12-Wk Vol", "Growth %", "Representative Query"]
    col_widths = [Inches(1.3), Inches(1.3), Inches(0.8), Inches(0.8), Inches(2.6)]

    for idx, (cell, text) in enumerate(zip(hdr_cells, headers)):
        cell.width = col_widths[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        # Background shading
        shading = parse_xml(r'<w:shd {} w:fill="2B3A67"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)

    for item in data["high_volume_gaps"]:
        row_cells = table_hv.add_row().cells
        vals = [
            f"{item['crop']}\n({item['state']})",
            item['domain'],
            str(item['total_volume']),
            f"{'+' if item['growth_pct']>0 else ''}{item['growth_pct']}%",
            f"\"{item['representative_question']}\""
        ]
        for idx, (cell, val) in enumerate(zip(row_cells, vals)):
            cell.width = col_widths[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(9)
            if idx == 3 and item['growth_pct'] >= 40:
                r.font.color.rgb = RGBColor(178, 65, 46)

    # Section 3: Fast-Growing Gaps
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    p_sec3 = doc.add_paragraph()
    style_heading(p_sec3, "3. Fast-Growing Knowledge Gaps (Emerging Spikes)", size_pt=14, color_rgb=(178, 65, 46))

    table_fg = doc.add_table(rows=1, cols=5)
    table_fg.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_fg.autofit = False

    hdr_cells_fg = table_fg.rows[0].cells
    for idx, (cell, text) in enumerate(zip(hdr_cells_fg, headers)):
        cell.width = col_widths[idx]
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(r'<w:shd {} w:fill="B2412E"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)

    for item in data["fast_growing_gaps"]:
        row_cells = table_fg.add_row().cells
        vals = [
            f"{item['crop']}\n({item['state']})",
            item['domain'],
            str(item['total_volume']),
            f"+{item['growth_pct']}%",
            f"\"{item['representative_question']}\""
        ]
        for idx, (cell, val) in enumerate(zip(row_cells, vals)):
            cell.width = col_widths[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(9)
            if idx == 3:
                r.font.bold = True
                r.font.color.rgb = RGBColor(178, 65, 46)

    # Section 4: Outreach Planning Recommendations
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    p_sec4 = doc.add_paragraph()
    style_heading(p_sec4, "4. Recommended Field & Content Outreach Plan", size_pt=14, color_rgb=(43, 58, 103))

    table_out = doc.add_table(rows=1, cols=4)
    table_out.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_out.autofit = False

    headers_out = ["Urgency", "Target (Crop / Domain / State)", "Score", "Recommended Action Plan"]
    widths_out = [Inches(1.0), Inches(1.8), Inches(0.6), Inches(3.4)]

    for idx, (cell, text) in enumerate(zip(table_out.rows[0].cells, headers_out)):
        cell.width = widths_out[idx]
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(r'<w:shd {} w:fill="241F1A"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)

    for item in data["outreach_plan"]:
        row_cells = table_out.add_row().cells
        vals = [
            item["urgency"],
            f"{item['crop']} - {item['domain']}\nState: {item['state']}",
            str(item["priority_score"]),
            item["recommended_action"],
        ]
        for idx, (cell, val) in enumerate(zip(row_cells, vals)):
            cell.width = widths_out[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(9)
            if idx == 0 and item["urgency"] == "Immediate":
                r.font.bold = True
                r.font.color.rgb = RGBColor(178, 65, 46)

    output_doc_path = os.path.join(os.path.dirname(__file__), "Weekly_GDB_Gap_Report.docx")
    doc.save(output_doc_path)
    print(f"Generated GDB Gap Report Word Document: {output_doc_path}")

if __name__ == "__main__":
    generate_report()
